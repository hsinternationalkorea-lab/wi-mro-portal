# -*- coding: utf-8 -*-
"""WI MRO 검색 포탈 — 직원 사용 가이드 v1"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Wholesale Industry\AI Assistant\MRO\PriceIntel\portal\사용가이드_직원용_v1.docx"

WI_BLUE = RGBColor(0x00, 0x89, 0xD0)
WI_SLATE = RGBColor(0x2B, 0x3A, 0x4E)
GRAY = RGBColor(0x66, 0x66, 0x66)
WARN = RGBColor(0xC0, 0x39, 0x2B)
OK_GREEN = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), "맑은 고딕")
rfonts.set(qn("w:ascii"), "맑은 고딕")
rfonts.set(qn("w:hAnsi"), "맑은 고딕")

section = doc.sections[0]
section.top_margin = Cm(2); section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)


def H(text, size=18, color=WI_SLATE, before=14, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = color
    r.font.name = "맑은 고딕"
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def P(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold
    r.font.name = "맑은 고딕"
    if color:
        r.font.color.rgb = color


def B(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    r.font.name = "맑은 고딕"; r.font.size = Pt(11)


def step_box(num, title, body):
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(15)
    cell_n = table.cell(0, 0); cell_b = table.cell(0, 1)
    cell_n.width = Cm(2); cell_b.width = Cm(15)
    p = cell_n.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(str(num))
    r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = WI_BLUE
    r.font.name = "맑은 고딕"
    cell_n.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    tc_pr = cell_n._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "E8EEF4"); tc_pr.append(shd)

    cell_b.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p1 = cell_b.paragraphs[0]
    r1 = p1.add_run(title)
    r1.font.size = Pt(13); r1.font.bold = True; r1.font.color.rgb = WI_SLATE
    r1.font.name = "맑은 고딕"
    p2 = cell_b.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(11); r2.font.name = "맑은 고딕"

    for cell in [cell_n, cell_b]:
        tc_pr = cell._tc.get_or_add_tcPr()
        b = OxmlElement("w:tcBorders")
        for s in ["top", "left", "bottom", "right"]:
            bd = OxmlElement(f"w:{s}")
            bd.set(qn("w:val"), "single"); bd.set(qn("w:sz"), "4")
            bd.set(qn("w:color"), "CCCCCC")
            b.append(bd)
        tc_pr.append(b)
    doc.add_paragraph()


def T(headers, rows, widths=None, header_fill="0089D0"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if widths:
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        if widths:
            cell.width = Cm(widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "맑은 고딕"
        tc = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), header_fill); tc.append(shd)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            if widths:
                cell.width = Cm(widths[ci])
            p = cell.paragraphs[0]
            r = p.add_run(str(val) if val is not None else "")
            r.font.size = Pt(10); r.font.name = "맑은 고딕"
    for row in t.rows:
        for cell in row.cells:
            tc = cell._tc.get_or_add_tcPr()
            b = OxmlElement("w:tcBorders")
            for s in ["top", "left", "bottom", "right"]:
                bd = OxmlElement(f"w:{s}")
                bd.set(qn("w:val"), "single"); bd.set(qn("w:sz"), "4")
                bd.set(qn("w:color"), "BBBBBB")
                b.append(bd)
            tc.append(b)
    doc.add_paragraph()


# ============== 표지 ==============
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("WI MRO 검색 포탈\n사용 가이드")
tr.font.size = Pt(28); tr.font.bold = True
tr.font.color.rgb = WI_SLATE; tr.font.name = "맑은 고딕"

sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("산업재 통합 검색 · 견적 · 발주")
sr.font.size = Pt(13); sr.font.color.rgb = GRAY; sr.font.name = "맑은 고딕"

mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = mp.add_run(f"v1 · {date.today().strftime('%Y-%m-%d')} · 시험가동 직원용")
mr.font.size = Pt(11); mr.font.color.rgb = GRAY; mr.font.name = "맑은 고딕"

doc.add_paragraph()

# ============== 0. 한 장 요약 ==============
H("이 가이드 한 장 요약", size=18, color=WI_BLUE)
P("WI MRO 포탈은 약 20,000개 산업재를 한 곳에서 검색·견적 받을 수 있는 사내 도구입니다.", indent=0.3)
P("기본 흐름:", bold=True, indent=0.3)
B("검색창에 키워드 입력 (예: 안전화 270mm)")
B("결과에서 상품 카드 → 수량 입력 → [카트 담기]")
B("우측 상단 [장바구니] → 회사정보 입력 → [견적 요청 보내기]")
B("WI 영업팀(Sales@wholesale-k.com)이 1일 내 회신")

# ============== 1. 접속 ==============
H("1. 포탈 접속", size=18)

step_box(1, "브라우저 열기",
         "Chrome, Edge, 또는 Whale 권장. 익스플로러(IE)는 미지원.")
step_box(2, "주소 입력",
         "주소창에 다음 입력:\n\nhttp://localhost:8501  (운영자 PC에서 시험)\n또는 안내된 IP 주소")
step_box(3, "메인 화면 확인",
         "WI 로고 + 검색창 + 9개 카테고리 버튼이 보이면 성공.\n총 등록 SKU 수가 화면 하단에 표시됨.")

# ============== 2. 검색하기 ==============
H("2. 검색하기 — 3가지 방법", size=18)

H("2.1 키워드 검색 (가장 일반적)", size=14, color=WI_BLUE, before=8)
P("검색창에 한글 키워드 입력 후 Enter.", indent=0.3)
P("좋은 검색어 예시:", bold=True, indent=0.3)
B("\"안전화 270mm\" — 사이즈까지 명확")
B("\"3M 마스크\" — 브랜드 + 종류")
B("\"드라이버\" — 일반명")
B("\"콤비네이션 렌치 17mm\" — 정밀")
P("나쁜 검색어 예시:", bold=True, color=WARN, indent=0.3)
B("\"공구\" — 너무 광범위 (수만 건 노출)")
B("\"좋은 거\" — 의미 없음")

H("2.2 모델번호 검색 (정확)", size=14, color=WI_BLUE)
P("기존 거래에서 알고 있는 모델번호 직접 입력.", indent=0.3)
P("예: \"TR-632E\", \"125-1010\", \"PB212LH\"", indent=0.3)
P("정확 매칭 시 1순위로 노출됨.", indent=0.3)

H("2.3 카테고리 둘러보기", size=14, color=WI_BLUE)
P("메인 화면 9개 카테고리 버튼 중 클릭:", indent=0.3)
T(["코드", "한글", "예시"],
  [
    ("M", "의료구급", "구급함, 일회용 마스크, 응급장비"),
    ("F", "체결부품", "볼트, 너트, 와셔, 체인"),
    ("L", "계측연구", "마이크로미터, 멀티미터, 측정기"),
    ("C", "크린용품", "방진복, 크린룸 장갑, 와이퍼"),
    ("S", "안전용품", "안전모, 안전화, 보호구, PPE"),
    ("E", "전기조명", "전구, 케이블타이, 콘센트"),
    ("P", "포장물류", "포장 비닐, 박스, 노끈, 캐비넷"),
    ("T", "공구류", "렌치, 드라이버, 망치, 드릴, 니퍼"),
    ("O", "사무일반", "USB, 바코드스캐너, 사무용품"),
  ],
  widths=[1.5, 3, 12])
P("카테고리 진입 후 좌측 사이드바에서 [브랜드 필터] 추가 가능.", indent=0.3)

# ============== 3. 결과 화면 사용법 ==============
H("3. 검색 결과 화면", size=18)
P("결과는 24개씩 4×6 그리드로 표시됩니다.", indent=0.3)
P("카드별 정보:", bold=True, indent=0.3)
B("이미지 (출처 사이트 hot-link)")
B("상품명, 규격, 제조사")
B("표시가 (단위 포함, 예: ₩89,000 / EA)")
B("WI 품번 (우리 내부 코드)")
B("DIRECT 뱃지 — 우리 직거래 SKU (우선순위 높음)")
P("좌측 사이드바:", bold=True, indent=0.3)
B("분류 (9개 카테고리, 하나 선택)")
B("브랜드 필터 (현재 결과의 TOP 10 + 카운트)")
P("상단 정렬:", bold=True, indent=0.3)
B("관련도 (기본) — 직거래·신뢰도 우선")
B("가격↑ — 저렴한 순")
B("가격↓ — 비싼 순")
B("신상품 — 등록 최근 순")
B("WI품번 — 코드 순")
P("페이지네이션:", bold=True, indent=0.3)
B("[← 이전] [페이지 X/N · 총 N건] [다음 →]")

# ============== 4. 장바구니 사용법 ==============
H("4. 장바구니 — 여러 품목 한 번에 견적", size=18, color=OK_GREEN)

step_box(1, "상품 발견 → 수량 입력 → [카트 담기]",
         "상품 카드 하단에 [수량] 입력칸과 [카트 담기] 버튼.\n수량 기본값 1, 원하는 수량으로 조정 후 클릭.")
step_box(2, "다른 상품 또 담기",
         "검색창에 다음 상품 입력 → 같은 방식으로 카트 담기.\n우측 상단 [장바구니 (N)] 카운트가 늘어남 (파란색 강조).")
step_box(3, "[장바구니] 클릭",
         "우측 상단 [장바구니 (N)] 클릭 → 카트 페이지 진입.\n담은 상품 라인별로 표시됨.")
step_box(4, "수량 조정 / 삭제",
         "각 라인의 수량 input으로 조정 가능.\n잘못 담은 건 [삭제] 버튼.")
step_box(5, "총 합계 확인",
         "하단에 \"총 N품목 · M개 · ₩X,XXX,XXX\" 표시.\n표시가 기준 합계 (할인·협상 후 변동 가능).")
step_box(6, "회사정보 입력 → [견적 요청 보내기]",
         "필수: 회사명, 담당자, 연락처, 이메일.\n선택: 사업자번호, 희망납기, 요청사항.\n[견적 요청 보내기] 클릭.")
step_box(7, "접수 번호 확인",
         "\"QR-260507-001 접수 완료\" 메시지.\n영업일 1일 내 회신 (Sales@wholesale-k.com).")

# ============== 5. 자주 하는 실수 ==============
H("5. 자주 하는 실수와 해결", size=18, color=WARN)
T(["증상", "원인", "해결"],
  [
    ("검색 결과가 너무 많음", "키워드가 너무 일반적", "사이즈·브랜드 같이 입력 (예: 안전화 → 안전화 270mm)"),
    ("결과가 0건", "오타 또는 등록 안된 SKU", "다른 키워드 시도, 또는 영업팀에 등록 요청"),
    ("이미지 안 보임", "출처 사이트 이미지 삭제", "정상 (일부 발생). 상품 정보는 그대로 사용"),
    ("카트 담아둔 게 사라짐", "브라우저 새로고침 또는 재시작", "다시 담기 (저장 기능은 향후 추가)"),
    ("수량 입력 시 1로 리셋", "Enter 안 누름", "수량 입력 후 마우스 클릭 또는 Tab"),
    ("견적 요청 후 회신 없음", "이메일 누락 또는 스팸함", "Sales@wholesale-k.com 확인 + 운영자 연락"),
  ],
  widths=[5, 5, 7])

# ============== 6. 시험가동 안내 ==============
H("6. 시험가동 기간 (5/7~5/13)", size=18)
P("이 시스템은 시험가동 단계입니다. 다음 사항에 협조 부탁드립니다.", indent=0.3)
B("실제 사용해본 후 발견되는 불편 즉시 보고 (홍기정 대표 또는 윤여준 매니저)")
B("발견된 오분류·오정보 캡처해서 공유 (WI품번 + 어떤 정보가 잘못)")
B("개선 아이디어 자유롭게 제안")
B("외부 고객에게는 시험가동 종료 후 공개 (5/14~)")
P("시험가동 동안:", bold=True, indent=0.3)
B("실제 견적 요청 → WI 영업팀이 정상 처리 (정식 견적서 발급)")
B("크레텍·석림랩텍 등 외부 카탈로그는 발주 시 우리가 대행")
B("WI 직거래 100 SKU는 기존 흐름 그대로")

# ============== 7. 문의 ==============
H("7. 문의처", size=18)
T(["내용", "연락처"],
  [
    ("시스템 사용 문의", "sale@wholesale-k.com (홍기정 대표)"),
    ("운영·견적 문의", "yjyoon@wholesale-k.com (윤여준 매니저)"),
    ("영업팀 공용", "Sales@wholesale-k.com"),
    ("긴급 (시스템 다운)", "010-XXXX-XXXX (운영자)"),
  ],
  widths=[5, 12])

# ============== 마무리 ==============
H("BUILT ON TRUST. DELIVERED WITH PRECISION.", size=14, color=WI_BLUE)
P("우리가 만드는 검색엔진의 시작입니다. 5일간의 시험가동에서 발견되는 모든 피드백이 다음 단계로 가는 발판입니다.", indent=0.3)
P("감사합니다.", bold=True, indent=0.3, color=WI_BLUE)

doc.save(OUT)
print(f"[OK] {OUT}")
