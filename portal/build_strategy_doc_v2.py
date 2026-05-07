# -*- coding: utf-8 -*-
"""
WI 그룹 AI 전략 방향 v2 — Group_Hub 3-tier 구조 반영
(_shared / WI / TeamS, 사업자번호 분리, 직거래 9개)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


WI_BLUE = RGBColor(0x00, 0x89, 0xD0)
WI_SLATE = RGBColor(0x2B, 0x3A, 0x4E)
WI_GRAY = RGBColor(0x6B, 0x72, 0x80)
WI_GREEN = RGBColor(0x38, 0x8E, 0x3C)
WI_ORANGE = RGBColor(0xF5, 0x7C, 0x00)


def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def set_run(run, size=10, bold=False, color=WI_SLATE, font="Pretendard"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_para(doc, text, size=10, bold=False, color=WI_SLATE, indent=0, space_after=4, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_section(doc, text, color=WI_SLATE):
    p = doc.add_paragraph()
    run = p.add_run("▎  ")
    set_run(run, size=13, bold=True, color=WI_BLUE)
    run = p.add_run(text)
    set_run(run, size=13, bold=True, color=color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)


def add_separator(doc, color="0089D0"):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def main():
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "WI_그룹_AI전략방향_v2.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ====== 헤더 ======
    add_para(doc, "Wholesale-K Group AI 전략 방향", size=22, bold=True, color=WI_SLATE, space_after=6)
    add_para(doc, "회사 두뇌(Corporate Brain) 구축 로드맵   |   v2.0  |  2026.05.07   |   대표이사 홍기정",
             size=10, color=WI_GRAY, space_after=6)
    add_separator(doc)

    # ====== 1. 핵심 메시지 ======
    add_section(doc, "1. 우리가 만드는 것")
    add_para(doc,
        "단순한 업무 자동화가 아닙니다. 회사의 모든 자산이 한 곳에 쌓이고, 정리되고, 검색되고, 증빙이 되는 — "
        "사람의 머리 대신 \"시스템\"이 회사 노하우를 보관하는 새로운 운영 모델입니다.",
        size=11, space_after=4)
    add_para(doc,
        "이것이 \"제3의 직원, AI 사업부\" — 회사 두뇌(Corporate Brain) 입니다.",
        size=11, bold=True, color=WI_BLUE, space_after=8)

    # ====== 2. 그룹 구조 — 3-tier ======
    add_section(doc, "2. 그룹 구조 (사업자 단위 분리)")
    structure = doc.add_table(rows=4, cols=3)
    structure.style = "Light Grid Accent 1"

    # 헤더
    headers_struct = ["영역", "사업자", "핵심"]
    for i, h in enumerate(headers_struct):
        cell = structure.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(cell, "0089D0")

    # 행
    rows_struct = [
        ("_shared", "그룹 공통", "메일·정책·교육·아카이브·그룹 재무 (연결손익)"),
        ("WI 유통 사업부", "사업자 A", "MRO 카탈로그·직거래 9개·견적·발주·이카운트 A"),
        ("TeamS 시공 사업부", "사업자 B", "방폭 전기공사·프로젝트·KOSHA 안전·이카운트 B"),
    ]
    for r, (a, b, c) in enumerate(rows_struct, 1):
        for ci, val in enumerate([a, b, c]):
            cell = structure.cell(r, ci)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run(run, size=10, bold=(ci == 0), color=WI_BLUE if ci == 0 else WI_SLATE)

    add_para(doc,
        "⚠ 사업자번호가 다르므로 ERP/HR/Finance는 사업부별 완전 분리. _shared 에는 그룹 차원 분석 데이터만.",
        size=9, color=WI_ORANGE, bold=True, space_after=8)

    # ====== 3. WI 직거래 9개 브랜드 ======
    add_section(doc, "3. WI 직거래 9개 브랜드 (핵심 차별화)")
    brand_table = doc.add_table(rows=3, cols=3)
    brand_table.style = "Light List"
    brands = [
        ("Fitok", "가스 피팅·밸브"),
        ("IEP", "정밀 임피던스"),
        ("Newson Gale", "정전기 grounding"),
        ("Halcarbon", "화학물질"),
        ("Uniphos", "가스 검지"),
        ("Twintec", "클린룸"),
        ("Valmet", "자동제어"),
        ("Atexon", "방폭"),
        ("TESI", "(분야 확정 예정)"),
    ]
    for i, (name, desc) in enumerate(brands):
        r, c = divmod(i, 3)
        cell = brand_table.cell(r, c)
        cell.text = ""
        p1 = cell.paragraphs[0]
        run = p1.add_run(name)
        set_run(run, size=10, bold=True, color=WI_BLUE)
        p2 = cell.add_paragraph()
        run2 = p2.add_run(desc)
        set_run(run2, size=9, color=WI_GRAY)
        set_cell_bg(cell, "F8F9FB")

    # ====== 4. 데이터 자산화 4가지 가치 ======
    add_section(doc, "4. 데이터 자산화의 4가지 가치")
    add_para(doc, "▸ 누적 (Accumulation)   정보가 흘러가지 않고 영구 축적", size=10)
    add_para(doc, "▸ 검색 (Retrieval)       필요할 때 0초로 답변 (RAG)", size=10)
    add_para(doc, "▸ 연결 (Connection)      임베딩 기반 패턴 발견 — 과거 사례가 미래 자산", size=10)
    add_para(doc, "▸ 증명 (Provability)     감사·분쟁·법적 증빙 즉시 제공", size=10, space_after=8)

    # ====== 5. 회사에 가져다줄 변화 ======
    add_section(doc, "5. 회사에 가져다줄 변화")
    add_para(doc, "▸ M&A 가치 ↑   데이터 자산화된 회사는 시장 평가에서 30~50% 프리미엄", size=10)
    add_para(doc, "▸ 사장 부재 리스크 분산   노하우가 사장 머리 → 시스템에 박힘", size=10)
    add_para(doc, "▸ 신입 교육 1/3 단축   RAG로 사내 지식 즉시 답변", size=10)
    add_para(doc, "▸ 거래처 신뢰 ↑   \"5초만에 자료 제공\" = 프로다운 회사", size=10)
    add_para(doc, "▸ 인건비 환산   AI 사업부 1대 = 영업·마케팅·운영·교육 11명 분의 자동화 효과", size=10)

    # ====== 페이지 2 ======
    doc.add_page_break()

    add_para(doc, "단계별 로드맵 — 12개월", size=18, bold=True, color=WI_SLATE, space_after=6)
    add_para(doc, "각 단계가 다음 단계의 토대가 됩니다", size=10, color=WI_GRAY, space_after=10)

    # ====== 6. 로드맵 ======
    add_section(doc, "6. 단계별 진행")
    rmap = doc.add_table(rows=5, cols=4)
    rmap.style = "Light Grid Accent 1"
    headers = ["단계", "시점", "핵심 작업", "산출물"]
    for i, h in enumerate(headers):
        cell = rmap.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run(run, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_cell_bg(cell, "0089D0")

    rows = [
        ("Phase 1", "5/8 ~", "MRO 통합 카탈로그 시험가동",
         "42,000+ SKU 검색·견적·관리자"),
        ("Phase 1.5", "5/15 ~", "유사상품 제안 + 멀티소스 dedup + 메일 IMAP 모니터링",
         "원가절감 자동 제안 / 견적 자동 분류"),
        ("Phase 2", "6/15 ~", "이카운트 ERP 통합 (사업자 A·B 별도) + 자동 보고서 + 마케팅 자동화",
         "포털↔ERP 양방향 sync, 위클리/거래처 보고서"),
        ("Phase 3", "9월 ~", "RAG 사내 지식 검색 + 신입 교육 챗봇 + AI 사업부",
         "전사 AI 협업, 모든 자산이 검색 가능한 두뇌"),
    ]
    for r, (a, b, c, d) in enumerate(rows, 1):
        for col, val in enumerate([a, b, c, d]):
            cell = rmap.cell(r, col)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            set_run(run, size=9, bold=(col == 0), color=WI_BLUE if col == 0 else WI_SLATE)

    # ====== 7. 인프라 ======
    add_section(doc, "7. 인프라 — 비용 대비 가치 극대화")
    add_para(doc, "▸ MacBook Pro M4 (사무실 24/7 가동)   AI Worker — 크롤링·자동화·코드 개발 거점", size=10)
    add_para(doc, "▸ 시놀로지 NAS (회사 기존 인프라)   Group_Hub/ 폴더 — 회사 모든 자산 영구 저장", size=10)
    add_para(doc, "▸ AWS (정식 호스팅)   거래처 접속 portal.koreaene.co.kr", size=10)
    add_para(doc, "▸ Supabase (DB)   상품·견적·거래처 클라우드 DB", size=10)
    add_para(doc, "▸ Windows 노트북 (사장 일상 사무)   메일·카톡·엑셀, 필요 시 Mac 원격 접속", size=10)
    add_para(doc, "▸ 메일 호스팅 (4개 도메인)   기존 그대로 — 자체 서버 안 함, AI Worker가 IMAP 통합 모니터링", size=10, space_after=8)

    # ====== 8. 직원과 함께 ======
    add_section(doc, "8. 직원과 함께 만들어갈 것")
    add_para(doc,
        "이 시스템은 사람을 대체하지 않습니다. 반복 업무를 AI에게 맡기고, 우리는 더 중요한 일에 집중하기 위함입니다.",
        size=10.5)
    add_para(doc, "", size=4)
    add_para(doc, "▸ 검색·견적·발주 같은 반복 업무 → AI Worker", size=10)
    add_para(doc, "▸ 거래처 관계, 새 사업 발굴, 전략 결정 → 사람", size=10)
    add_para(doc, "▸ 시험가동 기간(5/8 ~ 5/14) 직원 피드백이 시스템 방향을 결정합니다",
             size=10, bold=True, color=WI_BLUE)

    # ====== 9. 1년 후 ======
    add_section(doc, "9. 1년 후 우리 모습")
    add_para(doc,
        "사장이 일주일 자리를 비워도 회사가 돌아간다 — 이것이 진짜 회사 가치입니다.",
        size=11)
    add_para(doc,
        "우리는 \"사람이 운영하는 회사\"에서 \"데이터가 운영하는 시스템\"으로 진화합니다. "
        "500억 매출 목표 달성의 진짜 토대는 사람 수 늘리기가 아닌 시스템화입니다.",
        size=10, space_after=8)
    add_para(doc, "함께 만들어갑시다.", size=12, bold=True, color=WI_BLUE, space_after=8)

    # 푸터
    add_separator(doc, color="E5E8EC")
    add_para(doc,
        "㈜홀세일인더스트리   |   Built on Trust. Delivered with Precision.",
        size=8, color=WI_GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    doc.save(out_path)
    print(f"OK 생성 완료: {out_path}")


if __name__ == "__main__":
    main()
