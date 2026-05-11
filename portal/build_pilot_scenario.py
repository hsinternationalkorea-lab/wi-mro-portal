# -*- coding: utf-8 -*-
"""WI MRO 포탈 — 내부 시험가동 시나리오 v1"""
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Wholesale Industry\AI Assistant\MRO\PriceIntel\portal\시험가동_시나리오_v1.docx"

WI_BLUE = RGBColor(0x00, 0x89, 0xD0)
WI_SLATE = RGBColor(0x2B, 0x3A, 0x4E)
GRAY = RGBColor(0x66, 0x66, 0x66)
WARN = RGBColor(0xC0, 0x39, 0x2B)
OK_GREEN = RGBColor(0x2E, 0x7D, 0x32)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "맑은 고딕"
style.font.size = Pt(11)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
rfonts.set(qn("w:eastAsia"), "맑은 고딕")
rfonts.set(qn("w:ascii"), "맑은 고딕")
rfonts.set(qn("w:hAnsi"), "맑은 고딕")

section = doc.sections[0]
section.top_margin = Cm(2); section.bottom_margin = Cm(2)
section.left_margin = Cm(2.2); section.right_margin = Cm(2.2)


def H(text, size=18, color=WI_SLATE, before=14, after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = True; r.font.color.rgb = color
    r.font.name = "맑은 고딕"
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def P(text, bold=False, size=11, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    r = p.add_run(text)
    r.font.size = Pt(size); r.font.bold = bold
    r.font.name = "맑은 고딕"
    if color:
        r.font.color.rgb = color


def B(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
    r = p.runs[0] if p.runs else p.add_run()
    r.text = text
    r.font.name = "맑은 고딕"; r.font.size = Pt(11)


def T(headers, rows, widths=None, header_fill="0089D0"):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = False
    if widths:
        for i, w in enumerate(widths):
            t.columns[i].width = Cm(w)
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        if widths:
            cell.width = Cm(widths[i])
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.size = Pt(10); r.font.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        r.font.name = "맑은 고딕"
        tc = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), header_fill); tc.append(shd)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = t.cell(ri, ci)
            if widths:
                cell.width = Cm(widths[ci])
            p = cell.paragraphs[0]
            r = p.add_run(str(val) if val is not None else "")
            r.font.size = Pt(10); r.font.name = "맑은 고딕"
    for row in t.rows:
        for cell in row.cells:
            tc = cell._tc.get_or_add_tcPr()
            b = OxmlElement("w:tcBorders")
            for s in ["top", "left", "bottom", "right"]:
                bd = OxmlElement(f"w:{s}")
                bd.set(qn("w:val"), "single"); bd.set(qn("w:sz"), "4")
                bd.set(qn("w:color"), "BBBBBB")
                b.append(bd)
            tc.append(b)
    doc.add_paragraph()


def CHK(text, warn=False):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(2)
    box = p.add_run("☐  ")
    box.font.size = Pt(12); box.font.name = "맑은 고딕"
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.name = "맑은 고딕"
    if warn:
        r.font.color.rgb = WARN; r.font.bold = True


# ============== 표지 ==============
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = tp.add_run("WI MRO 검색 포탈\n시험가동 시나리오")
tr.font.size = Pt(28); tr.font.bold = True
tr.font.color.rgb = WI_SLATE; tr.font.name = "맑은 고딕"

sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sp.add_run("내부 직원 사용 검증 · 5일 운영")
sr.font.size = Pt(13); sr.font.color.rgb = GRAY; sr.font.name = "맑은 고딕"

mp = doc.add_paragraph(); mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = mp.add_run(f"v1 · {date.today().strftime('%Y-%m-%d')} · ㈜홀세일인더스트리")
mr.font.size = Pt(11); mr.font.color.rgb = GRAY; mr.font.name = "맑은 고딕"

doc.add_paragraph()

# ============== 1. 시험 목적 ==============
H("1. 시험가동 목적", size=18)
P("이번 시험가동은 '한국 B2B MRO 산업재 통합 검색 + 자동 거래 플랫폼' 비전을 향한 첫 검증 단계다. "
  "5일간 내부 직원이 실제 업무 흐름으로 사용해보고 다음을 평가한다.", indent=0.3)
P("핵심 검증 5가지", bold=True, indent=0.3)
B("검색 정확도 — 실제 자주 발주하는 SKU를 빠르고 정확히 찾는가")
B("UI 사용성 — 직원이 별도 교육 없이 직관적으로 사용 가능한가")
B("카탈로그 적합성 — 우리 업무에 필요한 SKU가 실제 노출되는가")
B("워크플로우 — 검색→장바구니→견적요청 흐름이 자연스러운가")
B("운영자 처리 — 들어온 견적 요청을 효과적으로 처리할 수 있는가")

# ============== 2. 시험 환경 ==============
H("2. 시험 환경", size=18)
T(["항목", "내용"],
  [
    ("URL", "http://localhost:8501 (시험기간) → koreaene.co.kr (정식)"),
    ("운영기간", f"{date.today().strftime('%Y-%m-%d')} ~ +5일"),
    ("등록 SKU", "약 17,500건 (WI 직거래 100 + 크레텍 17,400)"),
    ("Admin", "비밀번호 wi2026 (시험가동 임시)"),
    ("의뢰 수신", "Sales@wholesale-k.com (영업팀 공용)"),
    ("운영자", "윤여준 매니저 (yjyoon@wholesale-k.com)"),
  ],
  widths=[3.5, 13.5])

# ============== 3. 시험 참가자 ==============
H("3. 시험 참가자 (3~5명)", size=18)
P("아래 명단에서 3~5명을 선정해 시험에 투입.", indent=0.3)
T(["역할", "이름·연락처", "주요 작업", "기록"],
  [
    ("매니저(운영자)", "윤여준 / yjyoon@wholesale-k.com", "Admin 모드 의뢰 처리, 가격 협상", "처리 시간·이슈"),
    ("영업 1", "(추후 결정)", "검색·장바구니·견적 요청 (고객 시각)", "검색 성공률"),
    ("영업 2", "(추후 결정)", "카테고리 둘러보기·필터 사용", "사용성 점수"),
    ("현장 매니저 1", "(추후 결정)", "특정 SKU 검색 (시나리오 1~5)", "발견 시간"),
    ("관찰자 (대표)", "홍기정 / sale@wholesale-k.com", "전체 흐름 모니터링", "개선 의견"),
  ],
  widths=[3.5, 4.5, 5.5, 3.5])

# ============== 4. 시나리오 5개 ==============
H("4. 핵심 시나리오 5개", size=18)
P("각 시나리오는 실제 업무 상황을 모방. 시간·만족도·이슈를 기록.", indent=0.3)

H("4.1 시나리오 #1 — 단일 SKU 빠른 견적", size=14, color=WI_BLUE, before=12)
P("상황: 현장에서 갑자기 안전화 270mm 1켤레 필요. 5분 내 견적이 필요하다.", indent=0.3)
P("작업 순서:", bold=True, indent=0.3)
B("포탈 메인 화면 진입")
B("검색창에 '안전화 270mm' 입력")
B("결과 카드에서 적절한 상품 선택")
B("[수량 1] 입력 → [카트 담기]")
B("[장바구니] → 회사정보 입력 → [견적 요청 보내기]")
P("측정 항목:", bold=True, indent=0.3)
B("검색~견적요청 총 소요시간 (목표: 3분 이내)")
B("적절한 상품을 첫 화면에서 발견했는가 (목표: Yes)")
B("이미지·가격·단위가 명확한가 (1~5점)")

H("4.2 시나리오 #2 — 다중 SKU 일괄 견적", size=14, color=WI_BLUE)
P("상황: 신규 현장 셋업 — 안전모 5개, 작업장갑 10박스, 절단기 1대, 케이블타이 1봉 일괄 견적.", indent=0.3)
P("작업 순서:", bold=True, indent=0.3)
B("4가지 SKU를 각각 검색 + 카트 담기")
B("[장바구니] → 라인별 수량 확인 → 일괄 견적 요청")
P("측정 항목:", bold=True, indent=0.3)
B("4개 SKU 모두 검색 성공 (목표: 4/4)")
B("총 소요시간 (목표: 8분 이내)")
B("카트에서 수량 조정·삭제가 직관적인가")

H("4.3 시나리오 #3 — 모델번호 정확 검색", size=14, color=WI_BLUE)
P("상황: 기존 거래 SKU 재발주 — 정확한 모델번호 알고 있음 (예: TR-632E, BPK-HD).", indent=0.3)
P("작업 순서:", bold=True, indent=0.3)
B("검색창에 모델번호만 입력 (예: 'TR-632E')")
B("결과 1순위에 정확 매칭 노출되는지 확인")
P("측정 항목:", bold=True, indent=0.3)
B("정확 매칭 1순위 노출 (목표: 100%)")
B("WI 품번 자동 표시 확인")

H("4.4 시나리오 #4 — 카테고리 둘러보기", size=14, color=WI_BLUE)
P("상황: 어떤 안전 장비가 있는지 카탈로그를 둘러보고 싶다.", indent=0.3)
P("작업 순서:", bold=True, indent=0.3)
B("메인 화면에서 [안전용품] 카테고리 클릭")
B("좌측 사이드바 브랜드 필터 사용 (3M, 한세 등)")
B("정렬 [가격↑] 적용해서 저렴한 순으로 보기")
B("페이지 다음/이전 이동")
P("측정 항목:", bold=True, indent=0.3)
B("필터·정렬·페이지네이션 직관적 (1~5점)")
B("결과의 카테고리 분류 정확도 (오분류 발견 시 보고)")

H("4.5 시나리오 #5 — 운영자 의뢰 처리", size=14, color=WI_BLUE)
P("상황: 영업팀이 받은 견적 요청을 운영자(매니저)가 검토하고 처리한다.", indent=0.3)
P("작업 순서:", bold=True, indent=0.3)
B("우측 상단 [관리자 로그인] → wi2026")
B("[Admin] 의뢰 목록 페이지 진입")
B("새 견적 요청 클릭 → 라인별 원가·마진 확인")
B("크레텍에 발주 (수동) → 상태 'quoted'로 변경")
B("고객에게 견적서 회신 (Phase 1 견적서 시스템 활용)")
P("측정 항목:", bold=True, indent=0.3)
B("의뢰 검토 → 발주 → 회신 총 소요시간")
B("원가·마진 정보가 충분히 명확한가")
B("Phase 1 견적서와의 연계가 자연스러운가")

# ============== 5. 측정 항목 (5일 종합) ==============
H("5. 종합 측정 항목", size=18)
T(["항목", "목표", "측정 방법"],
  [
    ("검색 성공률", "85% 이상", "시도 100건 중 적절한 결과 노출"),
    ("평균 견적요청 시간", "5분 이내", "검색~카트~요청까지"),
    ("UI 만족도", "4점 이상 (5점)", "참가자 설문"),
    ("카테고리 정확도", "90% 이상", "랜덤 100개 SKU 검수"),
    ("이미지 노출률", "70% 이상", "전체 SKU 대비"),
    ("의뢰 처리 시간", "1일 내 회신", "운영자 측정"),
  ],
  widths=[5, 4, 7])

# ============== 6. 일자별 운영 계획 ==============
H("6. 5일 운영 계획", size=18)
T(["일자", "주요 활동", "참가자"],
  [
    ("D-Day (수)", "오리엔테이션 + 시나리오 #1 시도", "전원"),
    ("D+1 (목)", "시나리오 #2-3 + 자유 검색", "전원"),
    ("D+2 (금)", "시나리오 #4-5 + 운영자 처리", "전원"),
    ("D+3 (월)", "주말 정리 + 데이터 분석", "운영자·대표"),
    ("D+4 (화)", "회고 미팅 + 개선 우선순위 결정", "전원"),
  ],
  widths=[3, 9, 5])

# ============== 7. 의뢰 처리 흐름 ==============
H("7. 의뢰 처리 흐름 (운영자)", size=18)
P("시험가동 동안 의뢰가 들어오면 다음 절차로 처리한다.", indent=0.3)
B("Sales@wholesale-k.com 메일함 또는 Admin 의뢰 목록 확인")
B("의뢰 라인별 SKU 확인 — WI 품번 + 수량")
B("크레텍 SKU의 경우: 크레텍에 직접 발주 (오프라인)")
B("WI 직거래 SKU의 경우: 기존 공급사 발주")
B("Phase 1 견적서 시스템(WI_QuoteTemplate.xlsx)으로 견적서 작성")
B("고객에게 PDF 회신 (영업담당)")
B("Admin에서 status를 'quoted'로 변경")
B("발주 진행 → 'ordered' → 납품 완료 시 'completed'")

# ============== 8. 회고 ==============
H("8. 회고 미팅 (D+4)", size=18)
P("5일 운영 후 회고 미팅에서 다음 결정.", indent=0.3)
B("우선 개선 항목 TOP 3 (기능·UI·데이터 중)")
B("정식 출시 시기 (도메인 koreaene.co.kr)")
B("Phase 1+ 자동 발주 진입 시기")
B("추가 크롤링 대상 (석림랩텍·대한과학·나비엠알오)")
B("외부 고객 시범 운영 시기")

# ============== 9. 비상 대처 ==============
H("9. 비상 대처", size=18)
T(["증상", "원인", "조치"],
  [
    ("포탈 진입 안됨", "Streamlit 종료", "터미널에서 재실행 (운영자)"),
    ("이미지 안 보임", "출처 사이트 이미지 삭제", "정상 (일부 발생 가능)"),
    ("카테고리 잘못됨", "v3 매핑 한계", "메모해서 회고에서 수정"),
    ("의뢰 메일 안 옴", "Resend 설정 미완료", "Admin에서 직접 확인"),
    ("동시 사용자 많음", "무료 tier 한계", "유료 전환 검토 (Pro $25/월)"),
  ],
  widths=[4, 5, 7])

H("이 시험가동의 의미", size=14, color=OK_GREEN)
P("내일부터 5일은 우리 비전(2030 매출 500억 통합 플랫폼)의 첫 검증.", indent=0.3)
P("작은 시도지만, 데이터·사용자 피드백·운영 노하우가 진짜 무기가 된다.", indent=0.3)
P("BUILT ON TRUST. DELIVERED WITH PRECISION.", bold=True, color=WI_BLUE, indent=0.3)

doc.save(OUT)
print(f"[OK] {OUT}")
