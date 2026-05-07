# -*- coding: utf-8 -*-
"""
WI 그룹 AI 전략 방향 — 직원 배포용 워드 문서 (2장)
"""
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


WI_BLUE = RGBColor(0x00, 0x89, 0xD0)
WI_SLATE = RGBColor(0x2B, 0x3A, 0x4E)
WI_GRAY = RGBColor(0x6B, 0x72, 0x80)
WI_LIGHT = RGBColor(0xE5, 0xE8, 0xEC)
WI_GREEN = RGBColor(0x38, 0x8E, 0x3C)


def set_cell_bg(cell, color_hex):
    """셀 배경색 (HEX)"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_heading_blue(doc, text, size=22, color=WI_SLATE, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = "Pretendard"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
    run.font.size = Pt(size)
    run.font.bold = True
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subtitle(doc, text, size=11, color=WI_GRAY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Pretendard"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(12)
    return p


def add_body(doc, text, size=10, color=WI_SLATE, bold=False, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.name = "Pretendard"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_section_title(doc, text):
    p = doc.add_paragraph()
    # 좌측 파란 막대
    run = p.add_run("▎  ")
    run.font.color.rgb = WI_BLUE
    run.font.size = Pt(13)
    run.font.bold = True
    run = p.add_run(text)
    run.font.name = "Pretendard"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = WI_SLATE
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text, color=WI_SLATE):
    p = doc.add_paragraph(style="List Bullet")
    run = p.runs[0] if p.runs else p.add_run("")
    run.text = text
    run.font.name = "Pretendard"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
    run.font.size = Pt(10)
    run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35
    return p


def main():
    out_dir = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "WI_그룹_AI전략방향_v1.docx")

    doc = Document()

    # 페이지 여백
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # ====== 헤더 ======
    add_heading_blue(doc, "Wholesale-K Group AI 전략 방향", size=22, color=WI_SLATE)
    add_subtitle(doc, "회사 두뇌(Corporate Brain) 구축 로드맵   |   2026.05.07   |   대표이사 홍기정")

    # 구분선
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), '0089D0')
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ====== 1. 핵심 메시지 ======
    add_section_title(doc, "1. 우리가 만드는 것")
    add_body(doc,
        "단순한 업무 자동화가 아닙니다. "
        "회사의 모든 자산이 한 곳에 쌓이고, 정리되고, 검색되고, 증빙이 되는 — "
        "사람의 머리 대신 \"시스템\"이 회사 노하우를 보관하는 새로운 운영 모델입니다.",
        size=11, color=WI_SLATE)
    add_body(doc,
        "이것이 우리가 \"제3의 직원, AI 사업부\" 라고 부르는 회사 두뇌(Corporate Brain)입니다.",
        size=11, color=WI_BLUE, bold=True)

    # ====== 2. 통합 자산 영역 (표) ======
    add_section_title(doc, "2. 통합 관리 자산 (8개 영역)")
    table = doc.add_table(rows=3, cols=4)
    table.style = "Light Grid"
    items = [
        ("MRO 카탈로그", "포털·견적·발주"),
        ("보고서", "위클리·데일리·월간"),
        ("마케팅", "인스타·SNS·캠페인"),
        ("자매사 홈페이지", "팀에스·할로카본·WI"),
        ("회사 스탠다드", "SOP·정책·양식"),
        ("업무 프로세스", "영업·구매·운영·인사"),
        ("교육자료", "신입·제품·워크샵"),
        ("ERP·메일·아카이브", "이카운트·메일·과거자료"),
    ]
    for i, (title, desc) in enumerate(items):
        r, c = divmod(i, 4)
        cell = table.cell(r, c)
        cell.text = ""
        p1 = cell.paragraphs[0]
        run = p1.add_run(title)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WI_BLUE
        run.font.name = "Pretendard"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
        p2 = cell.add_paragraph()
        run2 = p2.add_run(desc)
        run2.font.size = Pt(9)
        run2.font.color.rgb = WI_GRAY
        run2.font.name = "Pretendard"
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
        set_cell_bg(cell, "F8F9FB")

    # ====== 3. 데이터 자산화 가치 4차원 ======
    add_section_title(doc, "3. 데이터 자산화의 4가지 가치")
    val_table = doc.add_table(rows=4, cols=2)
    val_table.style = "Light List"
    vals = [
        ("누적 (Accumulation)", "정보가 흘러가지 않고 영구 축적"),
        ("검색 (Retrieval)", "필요할 때 0초로 답변 (RAG)"),
        ("연결 (Connection)", "임베딩 기반 패턴 발견 — 과거 사례가 미래 자산"),
        ("증명 (Provability)", "감사·분쟁·법적 증빙 즉시 제공"),
    ]
    for i, (k, v) in enumerate(vals):
        c1 = val_table.cell(i, 0)
        c1.text = ""
        run = c1.paragraphs[0].add_run(k)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WI_BLUE
        run.font.name = "Pretendard"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
        c2 = val_table.cell(i, 1)
        c2.text = ""
        run2 = c2.paragraphs[0].add_run(v)
        run2.font.size = Pt(10)
        run2.font.color.rgb = WI_SLATE
        run2.font.name = "Pretendard"
        run2._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")

    add_body(doc, "", size=8)  # 여백

    # ====== 4. 비즈니스 임팩트 ======
    add_section_title(doc, "4. 회사에 가져다줄 변화")
    add_body(doc, "▸ M&A 가치 ↑   데이터 자산화된 회사는 시장 평가에서 30~50% 프리미엄", size=10)
    add_body(doc, "▸ 사장 부재 리스크 분산   노하우가 사장 머리 → 시스템에 박힘", size=10)
    add_body(doc, "▸ 신입 교육 기간 1/3 단축   RAG로 사내 지식 즉시 답변", size=10)
    add_body(doc, "▸ 거래처 신뢰 ↑   \"5초만에 자료 제공\" = 프로다운 회사", size=10)
    add_body(doc, "▸ 인건비 환산   AI 사업부 1대 = 영업·마케팅·운영·교육 11명 분의 자동화 효과", size=10)

    # ====== 페이지 2 시작 ======
    doc.add_page_break()

    add_heading_blue(doc, "단계별 로드맵 — 12개월", size=18, color=WI_SLATE)
    add_subtitle(doc, "각 단계가 다음 단계의 토대가 됩니다")

    # ====== 5. 로드맵 표 ======
    add_section_title(doc, "5. 단계별 진행")
    rmap = doc.add_table(rows=5, cols=4)
    rmap.style = "Light Grid Accent 1"
    headers = ["단계", "시점", "핵심 작업", "산출물"]
    for i, h in enumerate(headers):
        cell = rmap.cell(0, i)
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = "Pretendard"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")
        set_cell_bg(cell, "0089D0")

    rows = [
        ("Phase 1", "5/8 ~", "MRO 통합 카탈로그 시험가동", "42,000+ SKU 검색·견적·관리자"),
        ("Phase 1.5", "5/15 ~", "유사상품 제안 + 멀티소스 dedup + 메일 IMAP 모니터링", "원가절감 자동 제안 / 견적 자동 분류"),
        ("Phase 2", "6/15 ~", "이카운트 ERP 통합 + 자동 보고서 + 마케팅 자동화", "포털↔ERP 양방향 sync, 위클리/거래처 보고서"),
        ("Phase 3", "9월 ~", "RAG 사내 지식 검색 + 신입 교육 챗봇 + AI 사업부", "전사 AI 협업, 모든 자산이 검색 가능한 두뇌"),
    ]
    for r, (a, b, c, d) in enumerate(rows, 1):
        for col, val in enumerate([a, b, c, d]):
            cell = rmap.cell(r, col)
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            run.font.color.rgb = WI_SLATE if col != 0 else WI_BLUE
            run.font.bold = (col == 0)
            run.font.name = "Pretendard"
            run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")

    # ====== 6. 인프라 ======
    add_section_title(doc, "6. 인프라 — 비용 대비 가치 극대화")
    add_body(doc, "▸ Mac mini (사무실, 24/7 가동)   AI Worker — 크롤링·자동화·코드 개발 거점", size=10)
    add_body(doc, "▸ 시놀로지 NAS (회사 기존 인프라)   회사 모든 자산 영구 저장 (WI_Hub 폴더)", size=10)
    add_body(doc, "▸ AWS (정식 호스팅)   거래처가 접속하는 portal.koreaene.co.kr", size=10)
    add_body(doc, "▸ Supabase (DB)   상품·견적·거래처 클라우드 DB", size=10)
    add_body(doc, "▸ Windows 노트북 (사장 일상 사무)   메일·카톡·엑셀, 필요 시 Mac에 원격 접속", size=10)

    # ====== 7. 직원에게 ======
    add_section_title(doc, "7. 직원과 함께 만들어갈 것")
    add_body(doc,
        "이 시스템은 사람을 대체하지 않습니다. "
        "반복 업무를 AI에게 맡기고, 우리는 더 중요한 일에 집중하기 위함입니다.",
        size=10.5, color=WI_SLATE)
    add_body(doc, "", size=4)
    add_body(doc, "▸ 검색·견적·발주 같은 반복 업무 → AI Worker", size=10)
    add_body(doc, "▸ 거래처 관계, 새 사업 발굴, 전략 결정 → 사람", size=10)
    add_body(doc, "▸ 시험가동 기간(5/8 ~ 5/14) 직원 피드백이 시스템 방향을 결정합니다", size=10, bold=True, color=WI_BLUE)

    # ====== 8. 마무리 ======
    add_section_title(doc, "8. 1년 후 우리 모습")
    add_body(doc,
        "사장이 일주일 자리를 비워도 회사가 돌아간다 — "
        "이것이 진짜 회사 가치입니다.",
        size=11, color=WI_SLATE, bold=False)
    add_body(doc,
        "우리는 \"사람이 운영하는 회사\"에서 \"데이터가 운영하는 시스템\"으로 진화합니다. "
        "500억 매출 목표 달성의 진짜 토대는 사람 수 늘리기가 아닌 시스템화입니다.",
        size=10, color=WI_SLATE)
    add_body(doc, "", size=4)
    add_body(doc, "함께 만들어갑시다.", size=12, color=WI_BLUE, bold=True)

    # 푸터
    add_body(doc, "", size=4)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:color'), 'E5E8EC')
    pBdr.append(top)
    pPr.append(pBdr)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("㈜홀세일인더스트리   |   Built on Trust. Delivered with Precision.")
    run.font.size = Pt(8)
    run.font.color.rgb = WI_GRAY
    run.font.name = "Pretendard"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Pretendard")

    doc.save(out_path)
    print(f"OK 생성 완료: {out_path}")


if __name__ == "__main__":
    main()
